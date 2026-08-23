"""使用 python-docx 生成 .docx（排版更规范；未安装时回退到内置轻量生成器）。"""

import io

from .export import _contact_line, _sections


def build_docx_bytes(resume: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    b = resume.get("basic", {})

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(b.get("name") or "姓名")
    run.font.size = Pt(22)
    run.font.bold = True

    contact = _contact_line(b)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(contact)

    target = (b.get("target") or "").strip()
    if target:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.add_run(f"求职意向：{target}")

    for title, rows in _sections(resume):
        if not rows:
            continue
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.name = "微软雅黑"
            run.font.size = Pt(14)
        for head, bullets in rows:
            if head:
                hp = doc.add_paragraph()
                hr = hp.add_run(head)
                hr.font.bold = True
            for bl in bullets:
                doc.add_paragraph(bl, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
